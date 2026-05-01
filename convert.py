import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

with open('LIFEFORGE_Documentacion_Completa.md', 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

title = doc.add_heading('LifeForge - Documentación Completa', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Plataforma SaaS de Gestión de Hábitos y Nutrición')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

lines = content.split('\n')
in_table = False
current_table = None

for line in lines:
    line = line.strip()
    
    if not line:
        continue
    
    if line.startswith('# '):
        doc.add_heading(line[2:], 0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], 1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], 2)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], 3)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('| '):
        continue
    elif line.startswith('*') and not line.startswith('**'):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('```') or line.startswith('~~~'):
        continue
    else:
        clean_line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_line)
        clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)
        
        if clean_line and not clean_line.startswith('#'):
            p = doc.add_paragraph(clean_line)

doc.save('LIFEFORGE_Documentacion_Completa.docx')
print('Document created successfully!')